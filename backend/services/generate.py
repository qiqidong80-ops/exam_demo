import json
import io
import httpx
from fpdf import FPDF
from docx import Document
from config import DEEPSEEK_API_KEY, DEEPSEEK_BASE_URL, DEEPSEEK_MODEL

SYSTEM_PROMPT = """你是一位资深命题专家。根据给定的知识点和难度要求，生成 4 套试卷，每套含详细参考答案。

返回严格 JSON（不要 markdown 代码块）:
{
  "papers": [
    {
      "id": "paper-a",
      "title": "《课程名》模拟卷 A",
      "questions": [
        {"type": "选择题", "stem": "题目内容", "difficulty": "easy|medium|hard", "score": 5, "answer": "参考答案", "analysis": "解析"}
      ]
    }
  ]
}

要求:
1. 4 套试卷，每套 id 分别为 paper-a, paper-b, paper-c, paper-d
2. 每套 100 分，题型参考: 选择题(20分)、填空题(15分)、计算题(40分)、证明题(25分)
3. 难度配比: 严格按给定的 easy/medium/hard 分值比例分配
4. 覆盖所有给定的知识点，4 套合计覆盖更全面
5. 题目不得抄袭原题（换数据、换情境、换设问角度）
6. 每道题必须有 answer 和 analysis
7. 所有数学内容必须严格使用 LaTeX 格式：行内公式用 $...$ 包裹，独立公式用 $$...$$ 包裹；下标写作 x_1 或 x_{n+1}，上标写作 x^2 或 x^{n}，分数用 \\frac{}{}，根号用 \\sqrt{}，希腊字母用 \\pi \\alpha \\theta 等。禁止把 _ 和 ^ 等符号裸写在普通文本中（必须放在 $...$ 内）。
"""


async def run_generation(course: str, scope: str, kps: list[dict], difficulty: dict, ref_texts: list[str]) -> list[dict]:
    kp_summary = "\n".join([f"- {kp['name']} ({kp['difficulty']})" for kp in kps])
    ref_summary = "\n\n".join(ref_texts)[:30000]

    async with httpx.AsyncClient(timeout=600) as client:
        resp = await client.post(
            f"{DEEPSEEK_BASE_URL}/chat/completions",
            headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
            json={
                "model": DEEPSEEK_MODEL,
                "messages": [
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": f"课程: {course}\n范围: {scope or '全部'}\n难度比例: 基础{difficulty['easy']}% 中等{difficulty['medium']}% 难题{difficulty['hard']}%\n\n知识点:\n{kp_summary}\n\n往年真题参考:\n{ref_summary}"}
                ],
                "temperature": 0.7,
            }
        )
        resp.raise_for_status()
        data = resp.json()

    content = data["choices"][0]["message"]["content"].strip()
    if content.startswith("```"):
        content = content.split("\n", 1)[1].rsplit("\n", 1)[0]
    return json.loads(content).get("papers", [])


FONT_PATH = "C:/Windows/Fonts/simhei.ttf"

def export_pdf(paper: dict) -> bytes:
    pdf = FPDF()
    pdf.add_font("CJK", "", FONT_PATH, uni=True)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    # Title
    pdf.set_font("CJK", "", 14)
    pdf.cell(190, 10, paper["title"], align="C")
    pdf.ln(12)
    # Questions
    for i, q in enumerate(paper.get("questions", []), 1):
        pdf.set_font("CJK", "", 10)
        stem = f"{i}. [{q['type']}] ({q['score']}分) {q['stem']}"
        pdf.multi_cell(190, 7, stem)
        pdf.ln(2)
    # Answers header
    pdf.ln(4)
    pdf.set_font("CJK", "", 14)
    pdf.cell(190, 10, "参考答案", align="C")
    pdf.ln(10)
    # Answers
    for i, q in enumerate(paper.get("questions", []), 1):
        pdf.set_font("CJK", "", 10)
        pdf.multi_cell(190, 7, f"{i}. {q['answer']}")
        if q.get("analysis"):
            pdf.set_font("CJK", "", 9)
            pdf.multi_cell(190, 6, f"解析: {q['analysis']}")
            pdf.set_font("CJK", "", 10)
        pdf.ln(2)
    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def export_docx(paper: dict) -> bytes:
    doc = Document()
    doc.add_heading(paper["title"], 0)
    for i, q in enumerate(paper.get("questions", []), 1):
        doc.add_paragraph(f"{i}. [{q['type']}] ({q['score']}分) {q['stem']}")
    doc.add_page_break()
    doc.add_heading("参考答案", 1)
    for i, q in enumerate(paper.get("questions", []), 1):
        doc.add_paragraph(f"{i}. {q['answer']}")
        if q.get("analysis"):
            doc.add_paragraph(f"解析: {q['analysis']}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
